from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentModifier:
    def __init__(self, input_file, output_file):
        self.doc = Document(input_file)
        self.output_file = output_file
        self.new_mwh_value = "999.990000 MWh"
        self.new_digits = list("999990000")
        self.target_sections = {
            "Total production during period": False,
            "I‑REC(E) applied for": False,
            "I-REC(E) applied for": False
        }
        self.custom_fuel_code = "CUSTOM_ES"
        self.custom_fuel_description = "Custom Solar Fuel"
        self.custom_tech_code = "CUSTOM_TC"
        self.custom_tech_description = "Custom PV Tech"
        self.replacements = {
            "Facility ID/code": "TEST123456789",
            "Facility name": "Test Clean Power Ltd.",
        }
        self.date_replacements = {
            "Date": ("29", "03", "2025"),
            "Period start date": ("06", "05", "2027"),
            "Period end date": ("25", "06", "2027"),
        }

    def update_mwh_cell(self, cell):
        cell.text = ""
        paragraph = cell.add_paragraph(self.new_mwh_value)
        run = paragraph.runs[0]
        run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def find_energy_sources_table(self):
        for table_index, table in enumerate(self.doc.tables):
            for row in table.rows:
                if "Fuel" in row.cells[0].text and "Code(s)" in row.cells[1].text:
                    logger.info(f"Found matching table at index: {table_index}")
                    return table
        logger.error("Energy Sources table not found.")
        return None

    def update_energy_sources(self):
        table = self.find_energy_sources_table()
        if table:
            logger.info(f"Table has {len(table.rows)} rows.")

            for i, row in enumerate(table.rows):
                logger.info(f"\n[ROW {i}]")
                for j, cell in enumerate(row.cells):
                    logger.info(f"  Cell[{j}]: {cell.text.strip()}")

                if i == 3:
                    logger.info(f"\n[MODIFYING FUEL DATA - ROW {i}]")
                    logger.info(f"  Original Code: {row.cells[1].text}")
                    logger.info(f"  Original Description: {row.cells[2].text}")
                    row.cells[1].text = self.custom_fuel_code
                    row.cells[2].text = self.custom_fuel_description
                    logger.info(f"  -> Updated Code: {row.cells[1].text}")
                    logger.info(f"  -> Updated Description: {row.cells[2].text}")

                elif i == 5:
                    logger.info(f"\n[MODIFYING TECHNOLOGY DATA - ROW {i}]")
                    logger.info(f"  Original Code: {row.cells[1].text}")
                    logger.info(f"  Original Description: {row.cells[2].text}")
                    row.cells[1].text = self.custom_tech_code
                    row.cells[2].text = self.custom_tech_description
                    logger.info(f"  -> Updated Code: {row.cells[1].text}")
                    logger.info(f"  -> Updated Description: {row.cells[2].text}")
                else:
                    logger.info("[SKIPPED] No modification done on this row.")

    def update_production_details(self):
        current_section = None
        digit_count = 0

        for table in self.doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                
                for section in self.target_sections:
                    if section in row_text and not self.target_sections[section]:
                        current_section = section
                        self.target_sections[section] = True
                        digit_count = 0
                        break

                if not current_section:
                    continue

                for cell in row.cells:
                    text = cell.text.strip()
                    if "MWh" in text:
                        self.update_mwh_cell(cell)
                    elif text.isdigit() and digit_count < len(self.new_digits):
                        cell.text = self.new_digits[digit_count]
                        digit_count += 1

                if digit_count >= len(self.new_digits):
                    current_section = None

    def update_date_metadata(self):
        for table in self.doc.tables:
            for row in table.rows:
                for col_idx, cell in enumerate(row.cells):
                    label = cell.text.strip()
                    
                    if label in self.replacements and col_idx + 1 < len(row.cells):
                        row.cells[col_idx + 1].text = self.replacements[label]
                    
                    if label in self.date_replacements:
                        d, m, y = self.date_replacements[label]
                        try:
                            if label in ["Period start date", "Period end date"]:
                                row.cells[11].text = y
                                row.cells[7].text = m
                                row.cells[3].text = d
                            else:
                                row.cells[1].text = d
                                row.cells[2].text = m
                                row.cells[3].text = y
                            logger.info(f"Updated date for {label}: {d}-{m}-{y}")
                        except IndexError:
                            logger.error(f"Failed to update date for {label}")

    def process(self):
        self.update_energy_sources()
        self.update_production_details()
        self.update_date_metadata()
        self.doc.save(self.output_file)
        logger.info(f"Document saved as {self.output_file}")

def main():
    try:
        modifier = DocumentModifier(
            "SF-04-IssueRequest-v1.2 Example FIT devices.docx",
            "Modified_Data_Final.docx"
        )
        modifier.process()
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise

if __name__ == "__main__":
    main()
